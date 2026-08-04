# 公開できるサンプル文書(報告書.docx)を作る。中身はすべて架空。
#
#   .venv/bin/python sample/gen_report.py
#
# docx は python-docx で作る(SEKKEI「writer には橋を作らない」—
# その保存を writer が読めることは実物で確認済み)。
# サンプルは生成物 — 直すのはこのファイル。
import docx

d = docx.Document()

d.add_heading("月次業務報告(2026年7月)", level=1)

d.add_heading("概況", level=2)
d.add_paragraph(
    "7月の受注は3件(見積提出5件)。外壁塗装の引き合いが続いており、"
    "8月は足場資材の手配が山になる見込み。"
)

d.add_heading("受注の実績", level=2)
t = d.add_table(rows=4, cols=4)
for j, h in enumerate(["受注日", "件名", "発注元", "金額(税込)"]):
    t.rows[0].cells[j].text = h
for i, row in enumerate([
    ("7月3日", "外壁塗装工事", "株式会社みほん商事", "640,200円"),
    ("7月15日", "屋根の補修", "例示ビル管理", "231,000円"),
    ("7月28日", "駐車場の白線引き直し", "架空団地自治会", "88,000円"),
]):
    for j, v in enumerate(row):
        t.rows[i + 1].cells[j].text = v

d.add_heading("来月の予定", level=2)
d.add_paragraph("・8月上旬: みほん商事の現場に着手(足場は第1週に設置)")
d.add_paragraph("・8月中旬: 夏季休業(11日〜15日)")
d.add_paragraph("・8月下旬: 見積フォローの巡回")

p = d.add_paragraph()
r = p.add_run("以上")
r.bold = True

d.save("sample/報告書.docx")
print("書いた: sample/報告書.docx")
